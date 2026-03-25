from docx import Document
import os

doc = Document()
doc.add_heading("Luigi's Italian Restaurant - Menu", 0)

doc.add_heading("Appetizers", level=1)
doc.add_paragraph("1. Garlic Bread - $5.99\nFreshly baked bread with garlic butter and herbs.")
doc.add_paragraph("2. Bruschetta - $7.99\nToasted ciabatta topped with diced tomatoes, basil, and balsamic glaze.")

doc.add_heading("Main Courses", level=1)
doc.add_paragraph("1. Spaghetti Carbonara - $14.99\nClassic Roman pasta dish with eggs, cheese, pancetta, and black pepper.")
doc.add_paragraph("2. Margherita Pizza - $12.99\nTraditional pizza with San Marzano tomato sauce, fresh mozzarella, and basil.")
doc.add_paragraph("3. Chicken Parmesan - $16.99\nBreaded chicken breast topped with marinara and mozzarella, served over spaghetti.")

doc.add_heading("Desserts", level=1)
doc.add_paragraph("1. Tiramisu - $6.99\nCoffee-flavored Italian dessert with mascarpone cheese.")
doc.add_paragraph("2. Cannoli - $4.99\nPastry shells filled with sweet ricotta chocolate chip cream.")

output_path = os.path.abspath(os.path.join(os.path.dirname(__file__), "..", "docs", "menu.docx"))
doc.save(output_path)
print(f"✅ Saved fake menu DOCX in {output_path}")
